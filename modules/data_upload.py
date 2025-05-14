import requests
import streamlit as st
import pandas as pd
from docx import Document
import pdfplumber  # для таблиць з PDF
from io import StringIO

# Функція для зчитування DOCX файлів
def read_docx(file):
    doc = Document(file)
    dataframes = []

    # Зчитування таблиць
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        df = pd.DataFrame(rows)
        dataframes.append(df)

    # Якщо таблиць немає – зчитати текст
    if not dataframes:
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])
        return pd.DataFrame({"Текст": text.split("\n")})
    else:
        return pd.concat(dataframes, ignore_index=True)

# Функція для зчитування PDF файлів
def read_pdf(file):
    dataframes = []

    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                df = pd.DataFrame(table)
                dataframes.append(df)

    if not dataframes:
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return pd.DataFrame({"Текст": text.split("\n")})
    else:
        return pd.concat(dataframes, ignore_index=True)

# Головна функція додатку
def app():
    st.title("📁 Завантаження даних")

    # --- Завантаження з URL ---
    st.subheader("🌐 Завантаження з URL")
    url_input = st.text_input("Вставте URL до CSV, JSON або TXT-файлу:")

    if st.button("📥 Завантажити з URL") and url_input:
        try:
            response = requests.get(url_input)
            response.raise_for_status()
            content_type = response.headers.get("Content-Type", "")

            if "text/csv" in content_type or url_input.endswith(".csv"):
                df = pd.read_csv(StringIO(response.text))
            elif "application/json" in content_type or url_input.endswith(".json"):
                df = pd.read_json(StringIO(response.text))
            elif url_input.endswith(".txt"):
                df = pd.read_csv(StringIO(response.text), delimiter="\t")
            else:
                st.error("❌ Непідтримуваний формат або неправильний URL.")
                return

            st.success("✅ Дані з URL завантажено успішно!")
            st.dataframe(df)
            st.session_state["data"] = df

        except Exception as e:
            st.error(f"❌ Помилка при завантаженні з URL: {e}")

    # --- Локальне завантаження ---
    st.subheader("📂 Завантаження файлу з комп’ютера")
    uploaded_file = st.file_uploader(
        "Оберіть файл (CSV, Excel, JSON, TXT, Parquet, DOCX, PDF)", 
        type=["csv", "xlsx", "json", "txt", "parquet", "docx", "pdf"]
    )

    if uploaded_file:
        try:
            # Обробка локального файлу (залишається без змін)
            if uploaded_file.name.endswith(".csv"):
                df = pd.read_csv(uploaded_file)
            elif uploaded_file.name.endswith(".xlsx"):
                df = pd.read_excel(uploaded_file)
            elif uploaded_file.name.endswith(".json"):
                df = pd.read_json(uploaded_file)
            elif uploaded_file.name.endswith(".txt"):
                df = pd.read_csv(uploaded_file, delimiter="\t")
            elif uploaded_file.name.endswith(".parquet"):
                df = pd.read_parquet(uploaded_file)
            elif uploaded_file.name.endswith(".docx"):
                df = read_docx(uploaded_file)
            elif uploaded_file.name.endswith(".pdf"):
                df = read_pdf(uploaded_file)
            else:
                st.error("❌ Непідтримуваний формат файлу.")
                return

            st.success("✅ Дані успішно завантажено!")
            st.dataframe(df)
            st.session_state["data"] = df

        except Exception as e:
            st.error(f"❌ Помилка при завантаженні: {e}")
    else:
        st.info("⬆️ Завантажте файл або вкажіть URL для початку.")